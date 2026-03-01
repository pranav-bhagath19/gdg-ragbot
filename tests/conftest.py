"""
tests/conftest.py - Shared pytest fixtures.
Place this file at: rag-chatbot/tests/conftest.py
"""
import io
import os
import pytest
import pytest_asyncio
from pathlib import Path
from unittest.mock import patch, MagicMock

# Set test environment variables before imports
os.environ.setdefault("GROQ_API_KEY", "test-groq-key")
os.environ.setdefault("GEMINI_API_KEY", "test-gemini-key")
os.environ.setdefault("STORAGE_PATH", "/tmp/rag-test-data")
os.environ.setdefault("MAX_FILE_SIZE_MB", "10")
os.environ.setdefault("MAX_DOCS", "100")

from httpx import AsyncClient, ASGITransport
from fastapi.testclient import TestClient


@pytest.fixture(scope="session", autouse=True)
def setup_test_storage():
    """Create test storage directory."""
    Path("/tmp/rag-test-data").mkdir(parents=True, exist_ok=True)
    yield
    # Cleanup after all tests
    import shutil
    shutil.rmtree("/tmp/rag-test-data", ignore_errors=True)


@pytest.fixture(scope="session")
def app():
    """Create FastAPI app with mocked model loading."""
    import numpy as np

    # Mock the embedding model to avoid downloading 90MB in tests
    mock_model = MagicMock()
    mock_model.encode.return_value = np.random.rand(1, 384).astype("float32")

    with patch("services.embed._model", mock_model):
        with patch("services.embed.load_model", return_value=mock_model):
            from main import app as fastapi_app
            # Initialize vector store for tests
            from services import vectorstore
            vectorstore.init_store("/tmp/rag-test-data")
            yield fastapi_app


@pytest.fixture
def client(app):
    """Sync test client."""
    return TestClient(app)


@pytest.fixture
def sample_pdf_bytes():
    """Generate a minimal valid PDF in memory."""
    # Minimal PDF content
    pdf_content = b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj
4 0 obj << /Length 44 >> stream
BT /F1 12 Tf 100 700 Td (Hello World Test PDF) Tj ET
endstream endobj
5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000360 00000 n
trailer << /Size 6 /Root 1 0 R >>
startxref
441
%%EOF"""
    return pdf_content


@pytest.fixture
def sample_docx_bytes():
    """Generate a minimal DOCX file in memory."""
    from docx import Document
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test paragraph for unit testing the RAG chatbot.")
    doc.add_heading("Section 1", 1)
    doc.add_paragraph("Section 1 content with important information.")
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.fixture
def sample_xlsx_bytes():
    """Generate a minimal XLSX file in memory."""
    import openpyxl
    buf = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TestSheet"
    ws.append(["Name", "Age", "City"])
    ws.append(["Alice", 30, "New York"])
    ws.append(["Bob", 25, "Los Angeles"])
    ws.append(["Charlie", 35, "Chicago"])
    wb.save(buf)
    buf.seek(0)
    return buf.read()
